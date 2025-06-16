from flask import Flask, render_template, request, redirect, url_for, send_file, session, flash, jsonify, make_response
from flask_wtf import FlaskForm
from wtforms import StringField, IntegerField, SubmitField
from wtforms.validators import DataRequired
from io import BytesIO
import calendar
from collections import defaultdict, OrderedDict
from num2words import num2words
import json
from werkzeug.utils import secure_filename
from reportlab.lib.pagesizes import A4
from reportlab.lib import colors
from reportlab.lib.units import cm
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
import io
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.pdfbase import pdfmetrics
import os
from datetime import datetime
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

app = Flask(__name__)
app.secret_key = 'your_secret_key'

class BillEntryForm(FlaskForm):
    invoice_no = StringField('Invoice/Voucher No.', validators=[DataRequired()])
    date = StringField('Date (dd-mm-yyyy)', validators=[DataRequired()])
    taka = IntegerField('Taka', validators=[DataRequired()])
    name_of_medicine = StringField('Name of Medicine', validators=[DataRequired()])
    submit = SubmitField('Add Entry')

def number_to_words(n):
    # Simple number to words for up to 99999999
    return num2words(n, to='cardinal', lang='en').replace(',', '').title()

def month_name(date_str):
    try:
        month = int(date_str.split('-')[1])
        return calendar.month_name[month]
    except:
        return "Unknown"

def group_bills_by_month(bills):
    grouped = OrderedDict()
    # Sort by year, month, day
    for entry in sorted(bills, key=lambda x: (x['date'].split('-')[2], x['date'].split('-')[1], x['date'].split('-')[0])):
        year = entry['date'].split('-')[2]
        month = entry['date'].split('-')[1]
        key = f"{year}-{month}"
        if key not in grouped:
            grouped[key] = []
        grouped[key].append(entry)
    return grouped

def month_year_name(key):
    year, month = key.split('-')
    return f"{calendar.month_name[int(month)]} {year}"

@app.route('/', methods=['GET', 'POST'])
def index():
    form = BillEntryForm()
    if 'bills' not in session:
        session['bills'] = []
    bills = session['bills']

    if request.method == 'POST' and 'person' in request.form:
        session['person'] = request.form.get('person', 'rashida')
    person = session.get('person', 'rashida')

    # Edit mode
    edit_entry = session.pop('edit_entry', None)
    edit_sl_no = request.form.get('edit_sl_no') or session.pop('edit_sl_no', None)
    if edit_entry:
        form.invoice_no.data = edit_entry['invoice_no']
        form.taka.data = edit_entry['taka']
        form.name_of_medicine.data = edit_entry['name_of_medicine']
        # Ensure date is in DD-MM-YYYY format
        date_parts = edit_entry['date'].split('-')
        if len(date_parts) == 3 and len(date_parts[2]) == 4:
            edit_date_value = f"{date_parts[0]}-{date_parts[1]}-{date_parts[2]}"
        else:
            edit_date_value = edit_entry['date']
    else:
        edit_date_value = ''

    if form.validate_on_submit():
        date_str = request.form.get('date')
        if date_str and '-' in date_str:
            date_formatted = date_str
        else:
            date_formatted = ''
        if edit_sl_no:
            for b in bills:
                if str(b['sl_no']) == str(edit_sl_no):
                    b['invoice_no'] = form.invoice_no.data
                    b['date'] = date_formatted
                    b['taka'] = form.taka.data
                    b['name_of_medicine'] = form.name_of_medicine.data
                    break
        else:
            entry = {
                'sl_no': len(bills) + 1,
                'invoice_no': form.invoice_no.data,
                'date': date_formatted,
                'taka': form.taka.data,
                'name_of_medicine': form.name_of_medicine.data
            }
            bills.append(entry)
        session['bills'] = bills
        return redirect(url_for('index'))

    month_subtotals = defaultdict(int)
    for entry in bills:
        year = entry['date'].split('-')[2]
        month = entry['date'].split('-')[1]
        key = f"{year}-{month}"
        month_subtotals[key] += entry['taka']
    grouped_bills = group_bills_by_month(bills)
    month_names = {k: month_year_name(k) for k in grouped_bills.keys()}

    # Assign display_sl_no according to display order
    display_sl_no = 1
    for key in grouped_bills:
        for entry in grouped_bills[key]:
            entry['display_sl_no'] = display_sl_no
            display_sl_no += 1

    # Calculate grand_total from grouped_bills (displayed rows only)
    grand_total = sum(entry['taka'] for key in grouped_bills for entry in grouped_bills[key])

    if person == 'rashida':
        display_name = 'Rashida Sultana'
        certification_line2 = '2. Certified that the amount claimed in the bill has actually been incurred for the treatment of myself.'
        signature_block = ['Rashida Sultana', 'Ex-Election Commissioner', 'Bangladesh Election Commission']
    else:
        display_name = 'Md. Moklasur Rohman, Husband of Rashida Sultana'
        certification_line2 = '2. Certified that Md. Moklasur Rohman is my husband and he is not an employed govt. servant under the Govt. of Bangladesh and he does not claim any benefit in his own right.'
        signature_block = ['Rashida Sultana', 'Ex-Election Commissioner', 'Bangladesh Election Commission']

    amount_words = number_to_words(grand_total)
    certification = [
        'E. Fees for consultation for proper diagnosis and treatment of the patient vide Receipts detailed below:',
        '1. Certified that the amount claimed in the bill has actually been incurred for the treatment of myself.',
        '2. Certified that Md. Moklasur Rohman is my husband and he is not an employed govt. servant under the Govt. of Bangladesh and he does not claim any benefit in his own right.'
    ]
    cert_extra = [
        f"(vii) Certified that the medicines, drugs etc. included in the vouchers detailed above for the total cost amounting of <b>Total Taka {grand_total}/-(Only {amount_words} Taka)</b> were prescribed by me and were essential for the recovery and restoration of the health of Rashida Sultana, Self and of those medicines, drugs etc. we are not a dietary nature.",
        "(viii) Certified that neither these medicines, drugs etc. nor their effective substitutes were available at the time in the hospital.",
        f"(ix) Certified that consultation fee specified in the vouchers detailed above, amounting Tk.{grand_total} were actually necessary for the proper diagnosis and the treatment of the patient.",
        "(x) Certified that Rashida Sultana, Self whose signature is given above was attended to by me.",
        f"(xi) Certified that government Servant was attended at his residence owing to the absence remoteness of a suitable hospital or the severity of the illness and the amount of the cost of similar treatment as referred in sub-rule (2) of the Rule 8 is <b>Total Taka {grand_total}/-(Only {amount_words} Taka)</b>.",
        "(xii) Certified that the treatment of the husband of Rashida Sultana, Election Commissioner is neither prenatal nor post-natal in nature."
    ]
    return render_template('index.html', form=form, bills=bills, month_subtotals=month_subtotals, month_names=month_names, grand_total=grand_total, amount_words=amount_words, display_name=display_name, certification=certification, signature_block=signature_block, person=person, grouped_bills=grouped_bills, edit_entry=edit_entry, edit_date_value=edit_date_value, cert_extra=cert_extra, edit_sl_no=edit_sl_no)

@app.route('/edit/<int:sl_no>', methods=['POST'])
def edit_entry(sl_no):
    bills = session.get('bills', [])
    entry = next((b for b in bills if b['sl_no'] == sl_no), None)
    if not entry:
        flash('Entry not found!')
        return redirect(url_for('index'))
    session['edit_sl_no'] = sl_no
    session['edit_entry'] = entry
    return redirect(url_for('index'))

@app.route('/delete/<int:sl_no>', methods=['POST'])
def delete_entry(sl_no):
    bills = session.get('bills', [])
    bills = [b for b in bills if b['sl_no'] != sl_no]
    # Re-assign sl_no
    for idx, b in enumerate(bills, 1):
        b['sl_no'] = idx
    session['bills'] = bills
    return redirect(url_for('index'))

@app.route('/reset')
def reset():
    session.pop('bills', None)
    session.pop('person', None)
    return redirect(url_for('index'))

@app.route('/download_pdf')
def download_pdf():
    bills = session.get('bills', [])
    month_subtotals = defaultdict(int)
    for entry in bills:
        year = entry['date'].split('-')[2]
        month = entry['date'].split('-')[1]
        key = f"{year}-{month}"
        month_subtotals[key] += entry['taka']
    grand_total = sum(entry['taka'] for entry in bills)
    grouped_bills = group_bills_by_month(bills)
    month_names = {k: month_year_name(k) for k in grouped_bills.keys()}

    person = session.get('person', 'rashida')
    if person == 'rashida':
        display_name = 'Rashida Sultana'
        signature_block = ['Rashida Sultana', 'Ex-Election Commissioner', 'Bangladesh Election Commission']
    else:
        display_name = 'Md. Moklasur Rohman, Husband of Rashida Sultana'
        signature_block = ['Rashida Sultana', 'Ex-Election Commissioner', 'Bangladesh Election Commission']

    amount_words = number_to_words(grand_total)
    certification = [
        'E. Fees for consultation for proper diagnosis and treatment of the patient vide Receipts detailed below:',
        '1. Certified that the amount claimed in the bill has actually been incurred for the treatment of myself.',
        '2. Certified that Md. Moklasur Rohman is my husband and he is not an employed govt. servant under the Govt. of Bangladesh and he does not claim any benefit in his own right.'
    ]
    cert_extra = [
        f"(vii) Certified that the medicines, drugs etc. included in the vouchers detailed above for the total cost amounting of <b>Total Taka {grand_total}/-(Only {amount_words} Taka)</b> were prescribed by me and were essential for the recovery and restoration of the health of Rashida Sultana, Self and of those medicines, drugs etc. we are not a dietary nature.",
        "(viii) Certified that neither these medicines, drugs etc. nor their effective substitutes were available at the time in the hospital.",
        f"(ix) Certified that consultation fee specified in the vouchers detailed above, amounting Tk.{grand_total} were actually necessary for the proper diagnosis and the treatment of the patient.",
        "(x) Certified that Rashida Sultana, Self whose signature is given above was attended to by me.",
        f"(xi) Certified that government Servant was attended at his residence owing to the absence remoteness of a suitable hospital or the severity of the illness and the amount of the cost of similar treatment as referred in sub-rule (2) of the Rule 8 is <b>Total Taka {grand_total}/-(Only {amount_words} Taka)</b>.",
        "(xii) Certified that the treatment of the husband of Rashida Sultana, Election Commissioner is neither prenatal nor post-natal in nature."
    ]

    # Register Times New Roman font if available
    font_path = None
    for path in [
        '/Library/Fonts/Times New Roman.ttf',
        '/usr/share/fonts/truetype/msttcorefonts/Times_New_Roman.ttf',
        '/usr/share/fonts/truetype/msttcorefonts/Times_New_Roman.ttf',
        '/usr/share/fonts/truetype/msttcorefonts/Times_New_Roman.ttf',
        '/usr/share/fonts/truetype/msttcorefonts/Times New Roman.ttf',
        '/usr/share/fonts/truetype/msttcorefonts/times.ttf',
        os.path.expanduser('~/Library/Fonts/Times New Roman.ttf'),
    ]:
        if os.path.exists(path):
            font_path = path
            break
    if font_path:
        pdfmetrics.registerFont(TTFont('TimesNewRoman', font_path))
        base_font = 'TimesNewRoman'
    else:
        base_font = 'Times-Roman'
    styles = getSampleStyleSheet()
    for style_name in styles.byName:
        styles.byName[style_name].fontName = base_font
        styles.byName[style_name].fontSize = 12
    styles.add(ParagraphStyle(name='TableHeader', fontName=base_font, fontSize=12, alignment=1, spaceAfter=4, spaceBefore=4, leading=14, fontWeight='bold'))
    styles.add(ParagraphStyle(name='TableCell', fontName=base_font, fontSize=12, alignment=0, leading=14))
    styles.add(ParagraphStyle(name='RightCell', fontName=base_font, fontSize=12, alignment=2, leading=14))
    styles.add(ParagraphStyle(name='LeftCell', fontName=base_font, fontSize=12, alignment=0, leading=14))
    # Add bold style for the main title
    styles.add(ParagraphStyle(name='TitleBold', parent=styles['Title'], fontName=base_font+'-Bold' if 'TimesNewRoman' in base_font else 'Times-Bold', fontSize=16, alignment=1))

    buffer = io.BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=A4, leftMargin=2.54*cm, rightMargin=2.54*cm, topMargin=2.54*cm, bottomMargin=2.54*cm)
    elements = []

    elements.append(Paragraph("MEDICAL BILL FORM", styles['TitleBold']))
    elements.append(Spacer(1, 8))
    elements.append(Paragraph(f"Medical Attendance Bill of <b>{display_name}</b>, Ex-Election Commissioner, Bangladesh Election Commission (Separate Bill should be submitted for each patient).", styles['Normal']))
    elements.append(Spacer(1, 12))

    # Table header
    table_data = [[
        Paragraph('SL No.', styles['TableHeader']),
        Paragraph('Invoice/Voucher No.', styles['TableHeader']),
        Paragraph('Date', styles['TableHeader']),
        Paragraph('Taka', styles['TableHeader']),
        Paragraph('Name of Medicine', styles['TableHeader'])
    ]]
    sl_no = 1
    for key, entries in grouped_bills.items():
        for entry in entries:
            row = [
                Paragraph(str(sl_no), styles['TableCell']),
                Paragraph(str(entry['invoice_no']), styles['TableCell']),
                Paragraph(str(entry['date']), styles['TableCell']),
                Paragraph(str(entry['taka']), styles['TableCell']),
                Paragraph(str(entry['name_of_medicine']), styles['TableCell'])
            ]
            table_data.append(row)
            sl_no += 1
        # Add subtotal row
        subtotal_row = [
            "",  # SL No.
            "",  # Invoice/Voucher No.
            Paragraph(f"<b>SUBTOTAL ({month_names[key].upper()})</b>", styles['TableCell']),  # Date (bold)
            Paragraph(f"<b>{month_subtotals[key]:.2f}</b>", styles['TableCell']),  # Taka (bold)
            ""   # Name of Medicine
        ]
        table_data.append(subtotal_row)
    # Grand total row
    table_data.append([
        "",  # SL No.
        "",  # Invoice/Voucher No.
        Paragraph(f"<b>GRAND TOTAL</b>", styles['TableCell']),  # Date (bold)
        Paragraph(f"<b>{grand_total:.2f}</b>", styles['TableCell']),  # Taka (bold)
        ""   # Name of Medicine
    ])

    table = Table(table_data, colWidths=[2*cm, 4*cm, 3*cm, 3*cm, 5*cm])
    table.setStyle(TableStyle([
        ('BACKGROUND', (0,0), (-1,0), colors.lightgrey),
        ('GRID', (0,0), (-1,-1), 1, colors.black),
        ('ALIGN', (0,0), (-1,-1), 'CENTER'),
        ('FONTNAME', (0,0), (-1,-1), base_font),
        ('FONTSIZE', (0,0), (-1,-1), 12),
        ('BOTTOMPADDING', (0,0), (-1,0), 8),
        ('TOPPADDING', (0,0), (-1,0), 8),
        ('SPAN', (0,0), (0,0)),
    ]))
    elements.append(table)
    elements.append(Spacer(1, 16))
    elements.append(Paragraph(f"<b>Total Taka {grand_total}/-(Only {amount_words} Taka)</b>", styles['Normal']))
    elements.append(Spacer(1, 16))
    for line in certification:
        elements.append(Paragraph(line, styles['Normal']))
    elements.append(Spacer(1, 32))
    # Add extra space for signature before signature_block
    elements.append(Spacer(1, 40))  # Add space for signature
    for line in signature_block:
        elements.append(Paragraph(line, styles['RightCell']))
    elements.append(Spacer(1, 32))
    for line in cert_extra:
        elements.append(Paragraph(line, styles['Normal']))
    elements.append(Spacer(1, 32))
    # Add extra space for signature before medical attendant signature block
    elements.append(Spacer(1, 40))  # Add space for signature
    left_margin = int(0.55 * (A4[0] - 2 * 2.54 * cm))  # 55% of usable width
    for line in [
        "Signature of the Medical Attendant",
        "Registration No.",
        "Designation",
        "Post held"
    ]:
        elements.append(Paragraph(line, ParagraphStyle('Sig55', parent=styles['Normal'], leftIndent=left_margin)))
    doc.build(elements)
    buffer.seek(0)
    return send_file(buffer, as_attachment=True, download_name='medical_bill.pdf', mimetype='application/pdf')

@app.route('/download_docx')
def download_docx():
    bills = session.get('bills', [])
    month_subtotals = defaultdict(int)
    for entry in bills:
        year = entry['date'].split('-')[2]
        month = entry['date'].split('-')[1]
        key = f"{year}-{month}"
        month_subtotals[key] += entry['taka']
    grand_total = sum(entry['taka'] for entry in bills)
    grouped_bills = group_bills_by_month(bills)
    month_names = {k: month_year_name(k) for k in grouped_bills.keys()}

    person = session.get('person', 'rashida')
    if person == 'rashida':
        display_name = 'Rashida Sultana'
        signature_block = ['Rashida Sultana', 'Ex-Election Commissioner', 'Bangladesh Election Commission']
    else:
        display_name = 'Md. Moklasur Rohman, Husband of Rashida Sultana'
        signature_block = ['Rashida Sultana', 'Ex-Election Commissioner', 'Bangladesh Election Commission']

    amount_words = number_to_words(grand_total)

    # Create a new Document
    doc = Document()
    
    # Set margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(2.54)
        section.right_margin = Cm(2.54)

    # Add title
    title = doc.add_paragraph('MEDICAL BILL FORM')
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.runs[0].font.size = Pt(18)
    title.runs[0].font.bold = True
    doc.add_paragraph()  # Add space after title

    # Add header
    header = doc.add_paragraph(f'Medical Attendance Bill of {display_name}, Ex-Election Commissioner, Bangladesh Election Commission (Separate Bill should be submitted for each patient).')
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    header.runs[0].font.size = Pt(12)
    doc.add_paragraph()  # Add space after header

    # Create table
    table = doc.add_table(rows=1, cols=5)
    table.style = 'Table Grid'
    
    # Set column widths
    table.columns[0].width = Cm(2)  # SL No.
    table.columns[1].width = Cm(4)  # Invoice/Voucher No.
    table.columns[2].width = Cm(3)  # Date
    table.columns[3].width = Cm(3)  # Taka
    table.columns[4].width = Cm(5)  # Name of Medicine
    
    # Add header row
    header_cells = table.rows[0].cells
    headers = ['SL No.', 'Invoice/Voucher No.', 'Date', 'Taka', 'Name of Medicine']
    for i, header in enumerate(headers):
        header_cells[i].text = header
        header_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        header_cells[i].paragraphs[0].runs[0].font.bold = True
        header_cells[i].paragraphs[0].runs[0].font.size = Pt(12)

    # Add data rows
    sl_no = 1
    for key, entries in grouped_bills.items():
        for entry in entries:
            row_cells = table.add_row().cells
            row_cells[0].text = str(sl_no)
            row_cells[1].text = str(entry['invoice_no'])
            row_cells[2].text = str(entry['date'])
            row_cells[3].text = str(entry['taka'])
            row_cells[4].text = str(entry['name_of_medicine'])
            
            # Center align all cells and set font size
            for cell in row_cells:
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                cell.paragraphs[0].runs[0].font.size = Pt(11)
            
            sl_no += 1
        
        # Add subtotal row
        subtotal_row = table.add_row().cells
        subtotal_row[2].text = f'SUBTOTAL ({month_names[key].upper()})'
        subtotal_row[3].text = str(month_subtotals[key])
        subtotal_row[2].paragraphs[0].runs[0].font.bold = True
        subtotal_row[3].paragraphs[0].runs[0].font.bold = True
        subtotal_row[2].paragraphs[0].runs[0].font.size = Pt(11)
        subtotal_row[3].paragraphs[0].runs[0].font.size = Pt(11)
        
        # Center align subtotal cells
        for cell in subtotal_row:
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Add grand total row
    total_row = table.add_row().cells
    total_row[2].text = 'GRAND TOTAL'
    total_row[3].text = str(grand_total)
    total_row[2].paragraphs[0].runs[0].font.bold = True
    total_row[3].paragraphs[0].runs[0].font.bold = True
    total_row[2].paragraphs[0].runs[0].font.size = Pt(11)
    total_row[3].paragraphs[0].runs[0].font.size = Pt(11)
    
    # Center align total cells
    for cell in total_row:
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()  # Add space after table

    # Add total amount in words
    total_para = doc.add_paragraph(f'Total Taka {grand_total}/-(Only {amount_words} Taka)')
    total_para.runs[0].font.bold = True
    total_para.runs[0].font.size = Pt(12)
    doc.add_paragraph()  # Add space

    # Add certification
    cert_style = doc.styles.add_style('Certification', 1)
    cert_style.font.size = Pt(12)
    cert_style.font.name = 'Times New Roman'
    
    doc.add_paragraph('E. Fees for consultation for proper diagnosis and treatment of the patient vide Receipts detailed below:', style='Certification')
    doc.add_paragraph('1. Certified that the amount claimed in the bill has actually been incurred for the treatment of myself.', style='Certification')
    doc.add_paragraph('2. Certified that Md. Moklasur Rohman is my husband and he is not an employed govt. servant under the Govt. of Bangladesh and he does not claim any benefit in his own right.', style='Certification')

    # Add signature block
    doc.add_paragraph()  # Add space
    for line in signature_block:
        p = doc.add_paragraph(line)
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p.runs[0].font.size = Pt(12)

    # Add extra certifications
    doc.add_paragraph()
    cert_extra = [
        f"(vii) Certified that the medicines, drugs etc. included in the vouchers detailed above for the total cost amounting of Total Taka {grand_total}/-(Only {amount_words} Taka) were prescribed by me and were essential for the recovery and restoration of the health of Rashida Sultana, Self and of those medicines, drugs etc. we are not a dietary nature.",
        "(viii) Certified that neither these medicines, drugs etc. nor their effective substitutes were available at the time in the hospital.",
        f"(ix) Certified that consultation fee specified in the vouchers detailed above, amounting Tk.{grand_total} were actually necessary for the proper diagnosis and the treatment of the patient.",
        "(x) Certified that Rashida Sultana, Self whose signature is given above was attended to by me.",
        f"(xi) Certified that government Servant was attended at his residence owing to the absence remoteness of a suitable hospital or the severity of the illness and the amount of the cost of similar treatment as referred in sub-rule (2) of the Rule 8 is Total Taka {grand_total}/-(Only {amount_words} Taka).",
        "(xii) Certified that the treatment of the husband of Rashida Sultana, Election Commissioner is neither prenatal nor post-natal in nature."
    ]
    for cert in cert_extra:
        p = doc.add_paragraph(cert, style='Certification')

    # Add medical attendant signature block
    doc.add_paragraph()
    for line in [
        "Signature of the Medical Attendant",
        "Registration No.",
        "Designation",
        "Post held"
    ]:
        p = doc.add_paragraph(line)
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p.runs[0].font.size = Pt(12)

    # Save to BytesIO
    docx_bytes = BytesIO()
    doc.save(docx_bytes)
    docx_bytes.seek(0)

    return send_file(
        docx_bytes,
        as_attachment=True,
        download_name='medical_bill.docx',
        mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )

@app.route('/export')
def export_bills():
    bills = session.get('bills', [])
    person = session.get('person', 'rashida')
    today_str = datetime.now().strftime('%d-%m-%Y')
    if person == 'rashida':
        filename = f'Rashida_{today_str}.json'
    elif person == 'moklasur':
        filename = f'Moklasur_{today_str}.json'
    else:
        filename = f'bills_{today_str}.json'
    response = jsonify(bills)
    response.headers['Content-Disposition'] = f'attachment; filename={filename}'
    return response

@app.route('/import', methods=['POST'])
def import_bills():
    if 'file' not in request.files:
        flash('No file part')
        return redirect(url_for('index'))
    file = request.files['file']
    if file.filename == '':
        flash('No selected file')
        return redirect(url_for('index'))
    if file:
        try:
            data = json.load(file)
            # Validate data is a list of dicts with required keys
            if isinstance(data, list) and all(isinstance(x, dict) for x in data):
                for idx, b in enumerate(data, 1):
                    b['sl_no'] = idx
                session['bills'] = data
                flash('Bills imported successfully!')
            else:
                flash('Invalid file format!')
        except Exception as e:
            flash('Error importing file!')
    return redirect(url_for('index'))

if __name__ == '__main__':
    app.run(debug=True) 